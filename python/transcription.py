import os
import sys
import whisper
import subprocess
import math
import tempfile
from docx import Document
import soundfile as sf

if getattr(sys, "frozen", False):
    BASE_DIR = sys._MEIPASS  # PyInstaller
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

AUDIO_FOLDER = os.path.join(BASE_DIR, "audios")
os.makedirs(AUDIO_FOLDER, exist_ok=True)

FFMPEG_BIN = os.path.join(BASE_DIR, "bin", "ffmpeg.exe")

def extraer_audio(video_path, audio_path):
    cmd = [
        FFMPEG_BIN,
        "-i", video_path,
        "-vn",
        "-ac", "1",
        "-ar", "16000",
        "-c:a", "pcm_s16le",
        "-y",
        audio_path
    ]
    subprocess.run(cmd, check=True)

def transcribir_archivo(input_file, output_file):
    if not os.path.exists(input_file):
        raise FileNotFoundError(f"No se encontró el archivo: {input_file}")

    print("Progress: 5% - Iniciando transcripción...", flush=True)

    model = whisper.load_model("base")
    print("Progress: 15% - Modelo cargado correctamente.", flush=True)

    if input_file.lower().endswith((".mp4", ".mkv", ".mov", ".avi")):
        print("Progress: 25% - Extrayendo audio del video...", flush=True)
        audio_path = os.path.join(
            AUDIO_FOLDER, os.path.basename(input_file).rsplit(".", 1)[0] + ".wav"
        )
        extraer_audio(input_file, audio_path)
    else:
        audio_path = input_file

    print("Progress: 35% - Preparando audio...", flush=True)

    data, samplerate = sf.read(audio_path)
    duration = len(data) / samplerate
    segment_duration = 30  # segundos
    total_segments = math.ceil(duration / segment_duration)
    temp_dir = tempfile.mkdtemp()

    textos = []
    for i in range(total_segments):
        start = i * segment_duration
        end = min((i + 1) * segment_duration, duration)
        segment_data = data[int(start * samplerate):int(end * samplerate)]

        segment_path = os.path.join(temp_dir, f"segment_{i}.wav")
        sf.write(segment_path, segment_data, samplerate)

        progress = int(40 + (i / total_segments) * 55)
        print(f"Progress: {progress}% - Transcribiendo segmento {i + 1}/{total_segments}...", flush=True)

        result = model.transcribe(segment_path)
        textos.append(result["text"])

    texto = "\n".join(textos)

    print("Progress: 95% - Guardando archivo...", flush=True)
    if output_file.lower().endswith(".txt"):
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"Transcripción de {os.path.basename(input_file)}\n\n")
            f.write(texto)
    else:
        doc = Document()
        doc.add_heading(f"Transcripción de {os.path.basename(input_file)}", level=2)
        doc.add_paragraph(texto)
        doc.save(output_file)

    print("Progress: 100% - Completado ✅", flush=True)

    if audio_path != input_file and os.path.exists(audio_path):
        os.remove(audio_path)

    return f"✅ Transcripción completada en {output_file}"

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python transcription.py archivo_entrada salida.(docx|txt)")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        mensaje = transcribir_archivo(input_file, output_file)
        print(mensaje)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)







# import os
# import sys
# import whisper
# import subprocess
# from docx import Document

# if getattr(sys, "frozen", False):
#     BASE_DIR = sys._MEIPASS
# else:
#     BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# FFMPEG_BIN = os.path.join(BASE_DIR, "bin", "ffmpeg.exe")

# VIDEO_FOLDER = os.path.join(BASE_DIR, "videos")
# AUDIO_FOLDER = os.path.join(BASE_DIR, "audios")

# def extraer_audio(video_path, audio_path):
#     os.makedirs(AUDIO_FOLDER, exist_ok=True)

#     cmd = [
#         FFMPEG_BIN,
#         "-i", video_path,
#         "-vn",
#         "-ac", "1",
#         "-ar", "16000",
#         "-c:a", "pcm_s16le",
#         "-y",
#         audio_path
#     ]
#     subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

# def main(output_docx):
#     model = whisper.load_model("base")

#     doc = Document()

#     for file_name in os.listdir(VIDEO_FOLDER):
#         if file_name.lower().endswith((".mp4", ".mkv", ".mov", ".avi")):
#             video_path = os.path.join(VIDEO_FOLDER, file_name)
#             audio_path = os.path.join(AUDIO_FOLDER, file_name.rsplit(".", 1)[0] + ".wav")

#             extraer_audio(video_path, audio_path)
#             result = model.transcribe(audio_path)
#             doc.add_heading(f"Transcripción de {file_name}", level=2)
#             doc.add_paragraph(result["text"])

#     doc.save(output_docx)

# if __name__ == "__main__":
#     output_file = sys.argv[1] if len(sys.argv) > 1 else "transcripcion.docx"
#     main(output_file)



# import os
# import sys
# import whisper
# import subprocess
# from docx import Document

# BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# FFMPEG_BIN = os.path.join(BASE_DIR, "bin", "ffmpeg.exe")

# VIDEO_FOLDER = os.path.join(BASE_DIR, "videos")
# AUDIO_FOLDER = os.path.join(BASE_DIR, "audios")

# def extraer_audio(video_path, audio_path):
#     os.makedirs(AUDIO_FOLDER, exist_ok=True)

#     cmd = [
#         FFMPEG_BIN,
#         "-i", video_path,
#         "-vn",
#         "-ac", "1",
#         "-ar", "16000",
#         "-c:a", "pcm_s16le",
#         "-y",
#         audio_path
#     ]
#     subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

# def eliminar_archivo(path):
#     try:
#         if os.path.exists(path):
#             os.remove(path)
#     except Exception:
#         pass

# def limpieza_final():

#     if os.path.exists(AUDIO_FOLDER):
#         for f in os.listdir(AUDIO_FOLDER):
#             eliminar_archivo(os.path.join(AUDIO_FOLDER, f))

#     if os.path.exists(VIDEO_FOLDER):
#         for f in os.listdir(VIDEO_FOLDER):
#             eliminar_archivo(os.path.join(VIDEO_FOLDER, f))


# def main(output_docx):
#     model = whisper.load_model("base")
#     doc = Document()

#     for file_name in os.listdir(VIDEO_FOLDER):
#         if file_name.lower().endswith((".mp4", ".mkv", ".mov", ".avi")):
#             video_path = os.path.join(VIDEO_FOLDER, file_name)
#             audio_path = os.path.join(AUDIO_FOLDER, file_name.rsplit(".", 1)[0] + ".wav")

#             if extraer_audio(video_path, audio_path):
#                 try:
#                     result = model.transcribe(audio_path)
#                     doc.add_heading(f"Transcripción de {file_name}", level=2)
#                     doc.add_paragraph(result["text"])
#                 except Exception:
#                     pass

#     doc.save(output_docx)

#     limpieza_final()

# if __name__ == "__main__":
#     output_file = sys.argv[1] if len(sys.argv) > 1 else "transcripcion.docx"
#     main(output_file)



#  soporta videos y audios =>

# import os
# import sys
# import subprocess
# import whisper
# from docx import Document

# # === CONFIGURACIÓN ===
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# FFMPEG_BIN = os.path.join(BASE_DIR, "bin", "ffmpeg.exe")
# AUDIO_FOLDER = os.path.join(BASE_DIR, "audios")

# os.makedirs(AUDIO_FOLDER, exist_ok=True)

# # Extensiones
# VIDEO_EXT = (".mp4", ".mkv", ".mov", ".avi")
# AUDIO_EXT = (".mp3", ".wav", ".m4a", ".flac")


# def extraer_audio(video_path):
#     """
#     Convierte un video en un archivo wav temporal usando ffmpeg.
#     Retorna la ruta del wav generado o None si falla.
#     """
#     audio_name = os.path.splitext(os.path.basename(video_path))[0] + ".wav"
#     audio_path = os.path.join(AUDIO_FOLDER, audio_name)

#     cmd = [
#         FFMPEG_BIN,
#         "-i", video_path,
#         "-vn",
#         "-ac", "1",
#         "-ar", "16000",
#         "-c:a", "pcm_s16le",
#         "-y",
#         audio_path
#     ]

#     try:
#         subprocess.run(cmd, check=True)
#         return audio_path if os.path.exists(audio_path) else None
#     except subprocess.CalledProcessError:
#         return None


# def transcribir(model, file_path):
#     """
#     Transcribe un archivo de audio/video.
#     Retorna el texto o None si falla.
#     """
#     # Si es video → extraer audio
#     if file_path.lower().endswith(VIDEO_EXT):
#         audio_path = extraer_audio(file_path)
#         if not audio_path:
#             return None
#     elif file_path.lower().endswith(AUDIO_EXT):
#         audio_path = file_path  # usar directo
#     else:
#         return None

#     try:
#         result = model.transcribe(audio_path)
#         # Si era audio directo, no borrar
#         if file_path.lower().endswith(VIDEO_EXT):
#             os.remove(audio_path)  # limpiar wav temporal
#         return result["text"]
#     except Exception:
#         return None


# def main(output_docx, files):
#     """
#     Procesa una lista de archivos (videos o audios),
#     los transcribe y guarda en un DOCX.
#     """
#     model = whisper.load_model("base")
#     doc = Document()

#     for f in files:
#         text = transcribir(model, f)
#         if text:
#             doc.add_heading(f"Transcripción de {os.path.basename(f)}", level=2)
#             doc.add_paragraph(text)

#     doc.save(output_docx)


# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print("Uso: python transcription.py salida.docx archivo1 archivo2 ...")
#         sys.exit(1)

#     output_file = sys.argv[1]
#     input_files = sys.argv[2:]  # archivos pasados por línea de comando

#     main(output_file, input_files)
